"""Collect long-form text pulled out of Excel cells and write it to a single
companion Word document, with bookmarks that Excel hyperlinks can target.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

_BODY_FONT_SIZE = Pt(10.5)

_SENTENCE_BOUNDARY = re.compile(r"[。！？.!?]|;\s")
_SEMICOLON_SPLIT = re.compile(r";\s*")

_ASCII_FONT = "Times New Roman"
_EASTASIA_FONT = "游ゴシック"

SHEETS_EXCLUDED_FROM_WORD_EXPORT = {"Input_parameters"}


def _is_multi_sentence(text: str) -> bool:
    """Return True when text spans more than one sentence/segment.

    A single terminal punctuation mark at the very end of the string does
    not count as a boundary (e.g. "完了。" is one sentence). Only a boundary
    followed by further non-trivial content counts as a second sentence.
    """
    if not isinstance(text, str):
        return False
    stripped = text.strip()
    if not stripped:
        return False
    segments = [
        segment.strip()
        for segment in _SENTENCE_BOUNDARY.split(stripped)
        if segment.strip()
    ]
    return len(segments) > 1


@dataclass
class _WordExportItem:
    bookmark: str
    sheet_name: str
    column: str
    column_position: int
    row_index: int
    text: str

    @property
    def label(self) -> str:
        return f"{self.bookmark} — シート: {self.sheet_name} / 列: {self.column} (行 {self.row_index})"


@dataclass
class WordExportCollector:
    """Accumulates long-form cell text and assigns sequential bookmark IDs."""

    items: list[_WordExportItem] = field(default_factory=list)

    def add(self, sheet_name: str, column: str, column_position: int, row_index: int, text: str) -> str:
        bookmark = f"P{len(self.items) + 1}"
        self.items.append(_WordExportItem(bookmark, sheet_name, column, column_position, row_index, text))
        return bookmark

    def is_empty(self) -> bool:
        return not self.items


def _add_bookmark(paragraph, bookmark_name: str) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), "0")
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), "0")
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_internal_hyperlink(paragraph, text: str, anchor: str) -> None:
    """Add a run to `paragraph` that jumps to `anchor` within the same document."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    run_properties.append(style)
    run.append(run_properties)

    text_element = OxmlElement("w:t")
    text_element.set(qn("xml:space"), "preserve")
    text_element.text = text
    run.append(text_element)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _set_style_fonts(style, ascii_font: str, eastasia_font: str, size: Pt | None = None) -> None:
    """Set both the Latin (ascii/hAnsi) and Japanese (eastAsia) fonts on a style.

    python-docx's font.name only writes w:ascii/w:hAnsi; w:eastAsia must be
    set directly on the rFonts element or Japanese text keeps the Word
    theme default instead of the requested font.
    """
    style.font.name = ascii_font
    if size is not None:
        style.font.size = size
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), eastasia_font)


def _apply_document_fonts(document: Document) -> None:
    if "Normal" in document.styles:
        _set_style_fonts(document.styles["Normal"], _ASCII_FONT, _EASTASIA_FONT, _BODY_FONT_SIZE)
    for style_name in ("Title", "Heading 1"):
        if style_name in document.styles:
            _set_style_fonts(document.styles[style_name], _ASCII_FONT, _EASTASIA_FONT)


def _add_body_text(document: Document, text: str) -> None:
    """Write item body text, breaking "; "-joined lists onto separate lines."""
    segments = [segment.strip() for segment in _SEMICOLON_SPLIT.split(text) if segment.strip()]
    if len(segments) > 1:
        for segment in segments:
            document.add_paragraph(segment)
    else:
        document.add_paragraph(text)


def write_word_appendix(collector: WordExportCollector, output_path: str | Path) -> Path:
    """Write every collected item to one .docx: a table of contents page
    (with internal links) followed by one bookmarked section per item."""
    document = Document()
    _apply_document_fonts(document)

    document.add_heading("詳細補足資料", level=0)
    document.add_paragraph(
        "このドキュメントは、Excelレポート内のセルに収まらなかった長文の内容をまとめたものです。"
        "各項目はExcel側の該当セルからリンクされています。"
    )

    document.add_heading("目次", level=1)
    for item in collector.items:
        toc_paragraph = document.add_paragraph()
        _add_internal_hyperlink(toc_paragraph, item.label, item.bookmark)
    document.add_page_break()

    for item in collector.items:
        heading = document.add_heading(item.label, level=1)
        _add_bookmark(heading, item.bookmark)
        _add_body_text(document, item.text)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path